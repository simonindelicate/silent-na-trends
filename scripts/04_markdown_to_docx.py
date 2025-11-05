# scripts/04_markdown_to_docx.py
# Convert weekly_brief.md (markdown) → weekly_brief.docx
# - Clickable hyperlinks (explicit w:hyperlink + underline + blue)
# - Larger headings (H1/H2/H3)
# - Basic bullets / numbered lists

import argparse
import os
import re
from pathlib import Path
from typing import Optional
from docx import Document
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE

ROOT = Path(__file__).resolve().parents[1]
RUN_ID = os.getenv("RUN_ID")


def resolve_output_dir() -> Path:
    if RUN_ID:
        base = ROOT / "data" / "runs" / RUN_ID / "outputs"
    else:
        base = ROOT / "data" / "outputs"
    base.mkdir(parents=True, exist_ok=True)
    return base


OUT_DIR = resolve_output_dir()
LATEST_DIR = ROOT / "data" / "latest"
LATEST_DIR.mkdir(parents=True, exist_ok=True)

URL_RE = re.compile(r"(https?://[^\s)]+)")


def normalize_path(path: Path) -> Path:
    path = path.expanduser()
    return path if path.is_absolute() else (Path.cwd() / path)


def parse_args():
    parser = argparse.ArgumentParser(description="Convert weekly brief markdown to DOCX")
    parser.add_argument("--md", type=Path, help="Path to markdown file to convert")
    parser.add_argument("--docx", type=Path, help="Optional explicit DOCX output path")
    return parser.parse_args()


def discover_markdown(out_dir: Path, explicit: Optional[Path]) -> Path:
    if explicit:
        md_path = normalize_path(explicit)
        if not md_path.exists():
            raise FileNotFoundError(f"Markdown file not found: {md_path}")
        return md_path

    candidates = sorted(out_dir.glob("weekly_brief_*.md"))
    if candidates:
        return candidates[-1]

    fallback = out_dir / "weekly_brief.md"
    if fallback.exists():
        return fallback

    raise FileNotFoundError(
        f"No markdown brief found in {out_dir}. Run scripts/03_generate_report.py first."
    )


def resolve_docx_path(md_path: Path, explicit: Optional[Path]) -> Path:
    if explicit:
        doc_path = normalize_path(explicit)
        doc_path.parent.mkdir(parents=True, exist_ok=True)
        return doc_path

    if md_path.suffix:
        return md_path.with_suffix(".docx")
    return md_path.parent / f"{md_path.name}.docx"


def write_latest_docx(doc_path: Path, out_dir: Path, latest_dir: Path):
    data = doc_path.read_bytes()
    latest_path = out_dir / "weekly_brief.docx"
    latest_path.write_bytes(data)
    shared_latest = latest_dir / "weekly_brief.docx"
    shared_latest.write_bytes(data)

def style_document(doc: Document):
    # Base text
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    pf = normal.paragraph_format
    pf.space_after = Pt(6)
    pf.space_before = Pt(0)

    # Headings
    h1 = doc.styles["Heading 1"]
    h1.font.name = "Calibri Light"
    h1.font.size = Pt(26)
    h1.font.bold = True
    h1.paragraph_format.space_before = Pt(6)
    h1.paragraph_format.space_after = Pt(8)

    h2 = doc.styles["Heading 2"]
    h2.font.name = "Calibri"
    h2.font.size = Pt(20)
    h2.font.bold = True
    h2.paragraph_format.space_before = Pt(8)
    h2.paragraph_format.space_after = Pt(6)

    h3 = doc.styles["Heading 3"]
    h3.font.name = "Calibri"
    h3.font.size = Pt(14)
    h3.font.bold = True
    h3.paragraph_format.space_before = Pt(6)
    h3.paragraph_format.space_after = Pt(4)

def add_hyperlink(paragraph, url: str, text: str = None):
    """
    Insert a clickable external hyperlink into 'paragraph'.
    This does NOT depend on the 'Hyperlink' character style being present.
    """
    if text is None:
        text = url

    # 1) create relationship id to the external target
    part = paragraph.part
    r_id = part.relate_to(url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # 2) build the w:hyperlink wrapper with r:id
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    # 3) create a run with formatting (blue + underline)
    r = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")

    # underline
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    r_pr.append(u)

    # blue colour
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0000FF")
    r_pr.append(color)

    # (optional) no-proof to suppress spellcheck
    no_proof = OxmlElement("w:noProof")
    r_pr.append(no_proof)

    r.append(r_pr)

    # 4) text node
    t = OxmlElement("w:t")
    t.text = text
    r.append(t)

    # 5) append run to hyperlink, then hyperlink to paragraph
    hyperlink.append(r)
    paragraph._p.append(hyperlink)

def add_text_with_links(paragraph, text: str):
    """
    Append text to a paragraph, converting URLs into clickable hyperlinks.
    """
    pos = 0
    for m in URL_RE.finditer(text):
        before = text[pos:m.start()]
        if before:
            paragraph.add_run(before)
        url = m.group(1)
        add_hyperlink(paragraph, url, url)
        pos = m.end()
    tail = text[pos:]
    if tail:
        paragraph.add_run(tail)

def add_paragraph_with_links(doc: Document, text: str, style: str = None):
    p = doc.add_paragraph(style=style) if style else doc.add_paragraph()
    add_text_with_links(p, text)
    return p

def is_numbered(line: str) -> bool:
    # e.g., "1. Something", "12. Item"
    return re.match(r"^\s*\d+\.\s", line) is not None

def md_to_docx(md_text: str) -> Document:
    doc = Document()
    style_document(doc)

    bullet_mode = False
    for raw in md_text.splitlines():
        line = raw.rstrip()

        # Blank line → paragraph break
        if not line.strip():
            bullet_mode = False
            doc.add_paragraph("")
            continue

        # Headings
        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1); bullet_mode = False; continue
        if line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2); bullet_mode = False; continue
        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3); bullet_mode = False; continue

        # Bullets
        if line.lstrip().startswith("- "):
            if not bullet_mode:
                bullet_mode = True
            p = doc.add_paragraph(style=doc.styles["List Bullet"])
            add_text_with_links(p, line.lstrip()[2:].strip())
            continue

        # Numbered
        if is_numbered(line):
            content = re.sub(r"^\s*\d+\.\s", "", line).strip()
            p = doc.add_paragraph(style=doc.styles["List Number"])
            add_text_with_links(p, content)
            continue

        # Normal paragraph
        bullet_mode = False
        add_paragraph_with_links(doc, line)

    return doc

if __name__ == "__main__":
    args = parse_args()
    md_path = discover_markdown(OUT_DIR, args.md)
    md_text = md_path.read_text(encoding="utf-8")
    doc = md_to_docx(md_text)
    docx_path = resolve_docx_path(md_path, args.docx)
    doc.save(docx_path)
    write_latest_docx(docx_path, OUT_DIR, LATEST_DIR)
    print(f"Wrote DOCX -> {docx_path}")
