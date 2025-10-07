# scripts/03_generate_report.py
# Generate a weekly brief (Markdown + DOCX) from data/context/context.json using OpenAI.

import os
from pathlib import Path
import orjson
from dotenv import load_dotenv

load_dotenv()

OPENAI_KEY = os.getenv("OPENAI_API_KEY")
MODEL = os.getenv("OPENAI_MODEL", "gpt-5")  # set to a model your key has
ROOT = Path(__file__).resolve().parents[1]
CTX_PATH = ROOT / "data" / "context" / "context.json"
OUT_DIR = ROOT / "data" / "outputs"
OUT_DIR.mkdir(parents=True, exist_ok=True)

SYSTEM_PROMPT = (
    "You are a senior social strategist for a non-alcoholic beer brand serving a US audience. "
    "Write in warm, professional, approachable UK English (friendly, not stiff; never slangy). "
    "Be specific and evidence-led, but use smooth transitions and complete sentences. "
    "Avoid health claims and ‘sober shaming’. Use only the supplied context."
)

USER_PROMPT = """You are given structured context JSON (summary, top_posts, slang_candidates, reddit_posts, news_articles).

Produce a well-structured weekly report with these sections:

# Headline Summary (2 short paragraphs)
- What’s moving and why, in plain language. Name the biggest patterns.

## Top Trends (8 items)
For each:
- Title (friendly, descriptive)
- Why it matters (2–3 sentences; cite cross-platform presence or creator weight)
- 2–3 example links

## Slang & Phrases to Watch (10 items)
- Term — one-line gloss; add an example URL if present.

## Content Plan (5 themes)
For each theme:
- Rationale (2–3 sentences, grounded in the posts)
- Hook (10–12 words)
- Two formats (e.g., Reel, Carousel, TikTok) each with 3–6 beat bullets
- On-screen text (one line) + Caption (one sentence)
- 3–5 hashtags (blend brand + trend; UK English)
- Compliance notes (brief)

## Notables
- Product launches or notable creator posts (bulleted with links).

Tone: friendly, helpful, confident. Avoid clipped telegraph style."""

def load_context():
    if not CTX_PATH.exists():
        raise FileNotFoundError(f"Context not found: {CTX_PATH}. Run scripts/02_prepare_context.py first.")
    return orjson.loads(CTX_PATH.read_bytes())

def call_openai(system_text: str, user_text: str, context_json: dict, max_out_tokens=6000) -> str:
    if not OPENAI_KEY:
        raise RuntimeError("OPENAI_API_KEY is not set in .env")

    from openai import OpenAI
    from openai import BadRequestError

    client = OpenAI(api_key=OPENAI_KEY)
    ctx_str = orjson.dumps(context_json).decode("utf-8")
    messages = [
        {"role": "system", "content": system_text},
        {"role": "user", "content": user_text},
        {"role": "user", "content": f"Context JSON:\n```json\n{ctx_str}\n```"},
    ]

    # Prefer Responses API (GPT-5 family). Try max_output_tokens then max_completion_tokens.
    try:
        resp = client.responses.create(
            model=MODEL,
            input=messages,
            max_output_tokens=max_out_tokens,
        )
        return resp.output_text
    except TypeError:
        try:
            resp = client.responses.create(
                model=MODEL,
                input=messages,
                max_completion_tokens=max_out_tokens,
            )
            return resp.output_text
        except Exception:
            pass
    except BadRequestError as e:
        if "must use the chat.completions endpoint" not in str(e):
            raise

    # Fallback: Chat Completions (non-GPT-5 models)
    try:
        resp = client.chat.completions.create(
            model=MODEL,
            messages=messages,
            max_completion_tokens=max_out_tokens,
        )
        return resp.choices[0].message.content
    except BadRequestError as e:
        if "Use 'max_tokens' instead" not in str(e):
            raise
    resp = client.chat.completions.create(
        model=MODEL,
        messages=messages,
        max_tokens=max_out_tokens,
    )
    return resp.choices[0].message.content

def save_markdown(text: str, path: Path):
    path.write_text(text, encoding="utf-8")

def md_to_docx(md_text: str, out_path: Path):
    # Minimal markdown → docx (headings, bullets, paragraphs)
    from docx import Document
    from docx.shared import Pt
    doc = Document()
    doc.styles['Normal'].font.name = 'Calibri'
    doc.styles['Normal'].font.size = Pt(11)

    bullet_mode = False
    for raw in md_text.splitlines():
        line = raw.rstrip()
        if not line.strip():
            bullet_mode = False
            doc.add_paragraph("")
            continue
        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1); bullet_mode=False; continue
        if line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2); bullet_mode=False; continue
        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3); bullet_mode=False; continue
        if line.lstrip().startswith("- "):
            if not bullet_mode:
                bullet_mode = True
            p = doc.add_paragraph(line.lstrip()[2:].strip())
            p.style = doc.styles['List Bullet']
            continue
        # naive numbered list
        ls = line.lstrip()
        if ls[:2].isdigit() and ls[2:4] == ". ":
            p = doc.add_paragraph(ls[4:].strip())
            p.style = doc.styles['List Number']
            continue
        bullet_mode = False
        doc.add_paragraph(line)

    doc.save(out_path)

def main():
    ctx = load_context()
    md = call_openai(SYSTEM_PROMPT, USER_PROMPT, ctx, max_out_tokens=6000)
    md_path = OUT_DIR / "weekly_brief.md"
    save_markdown(md, md_path)
    docx_path = OUT_DIR / "weekly_brief.docx"
    md_to_docx(md, docx_path)
    print(f"Wrote {md_path}")
    print(f"Wrote {docx_path}")

if __name__ == "__main__":
    main()
